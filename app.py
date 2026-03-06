from flask import Flask, request, send_file
from docx import Document
import requests
import io
import re

app = Flask(__name__)

@app.route('/generate-docx', methods=['POST'])
def generate_docx():
    try:
        data = request.json
        lesson_data = data.get('lessonData', {})
        template_url = data.get('templateUrl')
        filename = data.get('filename', 'NTN_Lesson.docx')
        
        # Download template from Google Drive
        template_response = requests.get(template_url)
        template_file = io.BytesIO(template_response.content)
        
        # Open template
        doc = Document(template_file)
        
        # Helper function to strip markdown
        def strip_markdown(text):
            if not text:
                return ''
            # Remove markdown formatting
            text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)  # Bold italic
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)       # Bold
            text = re.sub(r'\*(.+?)\*', r'\1', text)           # Italic
            text = re.sub(r'###\s*', '', text)                 # Headers
            text = re.sub(r'^---+\s*$', '', text, flags=re.MULTILINE)  # Horizontal rules
            text = text.replace('●', '•')                      # Normalize bullets
            return text.strip()
        
        # Prepare replacements - strip markdown from all values
        replacements = {
            '{{GRADE_LEVEL}}': strip_markdown(lesson_data.get('gradeLevel', '')),
            '{{TOPIC}}': strip_markdown(lesson_data.get('topic', '')),
            '{{TIMEFRAME}}': strip_markdown(lesson_data.get('timeframe', '')),
            '{{STANDARD}}': strip_markdown(lesson_data.get('standard', '')),
            '{{OUTCOME}}': strip_markdown(lesson_data.get('outcome', '')),
            '{{RUBRIC_FOCUS}}': strip_markdown(lesson_data.get('rubricFocus', '')),
            '{{ACTIVITY_NAME}}': strip_markdown(lesson_data.get('activityName', '')),
            '{{SECTION_1_CONTENT}}': strip_markdown(lesson_data.get('section1Content', '')),
            '{{SECTION_1_TIME}}': strip_markdown(lesson_data.get('section1Time', '')),
            '{{SECTION_2_CONTENT}}': strip_markdown(lesson_data.get('section2Content', '')),
            '{{SECTION_2_TIME}}': strip_markdown(lesson_data.get('section2Time', '')),
            '{{SECTION_3_CONTENT}}': strip_markdown(lesson_data.get('section3Content', '')),
            '{{SECTION_3_TIME}}': strip_markdown(lesson_data.get('section3Time', '')),
            '{{SECTION_4_CONTENT}}': strip_markdown(lesson_data.get('section4Content', '')),
            '{{SECTION_4_TIME}}': strip_markdown(lesson_data.get('section4Time', '')),
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                for run in paragraph.runs:
                                    if key in run.text:
                                        run.text = run.text.replace(key, value)
        
        # Save to BytesIO
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        return {'error': str(e)}, 500

@app.route('/health', methods=['GET'])
def health():
    return {'status': 'healthy'}, 200

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080)
